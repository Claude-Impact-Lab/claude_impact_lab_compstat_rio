"""
Gerador do "Relatório Analítico de Área" em .docx — segue o template
do anexo do briefing (página 11 em diante).

Entrada:
    area:           dict do real.json (uma das entries em data.areas)
    reference_date: string ISO da data de referência do dataset

Saída:
    bytes do .docx serializado, prontos pra StreamingResponse.

O layout segue a estrutura consolidada do CompStat Municipal:
    1. Cabeçalho + área/período
    2. Resumo Executivo (perguntas norteadoras)
    3. Ocorrências Criminais (identificação, indicadores, distribuição, temporal)
    4. Dinâmica Criminal
    5. Efetivo Empregado — Força Municipal
    6. Fatores de Incidência Criminal
    7. Plano de Ação e Responsabilização
"""

from __future__ import annotations

import io
from collections import Counter, defaultdict
from typing import Any

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .map_renderer import render_map

# Paleta institucional (mesma do frontend)
COLOR_PRIMARY = RGBColor(0x1E, 0x3A, 0x5F)
COLOR_ACCENT = RGBColor(0x0E, 0xA5, 0xE9)
COLOR_MUTED = RGBColor(0x64, 0x74, 0x8B)
COLOR_RISK_CRITICAL = RGBColor(0xDC, 0x26, 0x26)
COLOR_RISK_HIGH = RGBColor(0xEA, 0x58, 0x0C)
COLOR_RISK_MEDIUM = RGBColor(0xD9, 0x77, 0x06)
SHADE_HEADER = "1E3A5F"   # azul-marinho
SHADE_SUBHEAD = "E2E8F0"  # cinza muito claro
SHADE_ROW_ALT = "F8FAFC"  # quase branco
SHADE_CRITICAL = "FEE2E2"
SHADE_HIGH = "FED7AA"
SHADE_MEDIUM = "FEF3C7"


# ---------------------------------------------------------------------------
# Low-level helpers
# ---------------------------------------------------------------------------

def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_table_borders(table, color: str = "94A3B8", sz: str = "4") -> None:
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), sz)
        b.set(qn("w:color"), color)
        borders.append(b)
    tbl_pr.append(borders)


def _set_run_color(run, color: RGBColor) -> None:
    run.font.color.rgb = color


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color: RGBColor | None = None,
    size: int | None = None,
    align: int | None = None,
) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if color is not None:
        _set_run_color(run, color)
    if size is not None:
        run.font.size = Pt(size)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    _set_run_color(run, COLOR_PRIMARY)
    run.font.size = Pt(13 if level == 1 else 11)


def _add_para(doc: Document, text: str, *, italic: bool = False,
              size: int = 10, color: RGBColor | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        _set_run_color(run, color)


def _header_row(table, labels: list[str]) -> None:
    """Aplica shading azul-marinho + texto branco bold na primeira linha."""
    for cell, label in zip(table.rows[0].cells, labels):
        _set_cell_shading(cell, SHADE_HEADER)
        _set_cell_text(cell, label, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _subheader_cell(cell, text: str) -> None:
    _set_cell_shading(cell, SHADE_SUBHEAD)
    _set_cell_text(cell, text, bold=True, size=10)


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _add_cover(doc: Document, area: dict, reference_date: str | None) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("RELATÓRIO ANALÍTICO DE ÁREA")
    run.bold = True
    run.font.size = Pt(18)
    _set_run_color(run, COLOR_PRIMARY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Subsídio para Reunião de CompStat")
    sub_run.font.size = Pt(11)
    sub_run.italic = True
    _set_run_color(sub_run, COLOR_MUTED)

    # Tabela: Área de análise | Período de análise
    table = doc.add_table(rows=2, cols=2)
    _set_table_borders(table)
    table.columns[0].width = Cm(8.5)
    table.columns[1].width = Cm(8.5)
    _subheader_cell(table.cell(0, 0), "Área de análise")
    _subheader_cell(table.cell(0, 1), "Período de análise")
    _set_cell_text(table.cell(1, 0), area["name"], bold=True, size=10)
    period = "Janela de 90 dias"
    if reference_date:
        period += f" · referência {reference_date}"
    _set_cell_text(table.cell(1, 1), period, size=10)
    doc.add_paragraph()  # spacer


def _add_map_image(doc: Document, area: dict) -> None:
    """Renderiza e embute o mapa de calor da área (polígono + heatmap + câmeras)."""
    try:
        png = render_map(area)
    except Exception as exc:  # noqa: BLE001
        _add_para(doc, f"Mapa indisponível: {exc}", italic=True, color=COLOR_MUTED, size=9)
        return
    if not png:
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    run.add_picture(io.BytesIO(png), width=Cm(17))
    _add_para(
        doc,
        "Mapa: polígono da Força Municipal (azul tracejado), heatmap das "
        "ocorrências dos últimos 90 dias e câmeras instaladas. Tile base: "
        "CartoDB Voyager.",
        italic=True, color=COLOR_MUTED, size=8,
    )


def _add_resumo_executivo(doc: Document, area: dict) -> None:
    _add_heading(doc, "RESUMO EXECUTIVO", level=1)
    items = area.get("executiveSummary") or []
    if not items:
        _add_para(doc, "Resumo executivo não gerado para esta área.", italic=True,
                  color=COLOR_MUTED)
        return
    table = doc.add_table(rows=len(items) + 1, cols=3)
    _set_table_borders(table)
    _header_row(table, ["Pergunta norteadora", "Diagnóstico com base nos dados",
                        "Fontes consultadas"])
    table.columns[0].width = Cm(5.5)
    table.columns[1].width = Cm(8.5)
    table.columns[2].width = Cm(3.0)
    for i, item in enumerate(items, start=1):
        _set_cell_text(table.cell(i, 0), item.get("q", ""), bold=True, size=10)
        _set_cell_text(table.cell(i, 1), item.get("a", ""), size=10)
        sources = item.get("sources") or []
        _set_cell_text(table.cell(i, 2), " · ".join(sources), size=9,
                       color=COLOR_MUTED)
        if i % 2 == 0:
            for c in table.rows[i].cells:
                _set_cell_shading(c, SHADE_ROW_ALT)


def _add_identificacao(doc: Document, area: dict) -> None:
    _add_heading(doc, "1. OCORRÊNCIAS CRIMINAIS", level=1)
    _add_heading(doc, "IDENTIFICAÇÃO DA ÁREA", level=2)
    rows = [
        ("Área FM", area["name"]),
        ("AISP", area.get("aisp", "—")),
        ("Bairro", area.get("bairro", "—")),
        ("Número de trechos críticos", str(len(area.get("coincidences") or []))),
        ("Câmeras instaladas no polígono", str(len(area.get("cameras") or []))),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    _set_table_borders(table)
    table.columns[0].width = Cm(5.5)
    table.columns[1].width = Cm(11.5)
    for i, (k, v) in enumerate(rows):
        _subheader_cell(table.cell(i, 0), k)
        _set_cell_text(table.cell(i, 1), v, size=10)


def _add_indicadores(doc: Document, area: dict) -> None:
    _add_heading(doc, "INDICADORES DO PERÍODO", level=2)
    kpis = area.get("kpis") or {}
    table = doc.add_table(rows=2, cols=5)
    _set_table_borders(table)
    _header_row(table, ["Janela", "Ocorrências (30d)", "Variação vs mês anterior",
                        "Fatores urbanos ativos", "Denúncias (90d)"])
    _set_cell_text(table.cell(1, 0), "30 dias",
                   align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _set_cell_text(table.cell(1, 1), str(kpis.get("ocorrencias_30d", "—")),
                   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    var = kpis.get("ocorrencias_var", 0)
    _set_cell_text(table.cell(1, 2), f"{var:+d}%",
                   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11,
                   color=(COLOR_RISK_CRITICAL if var > 5
                          else COLOR_RISK_MEDIUM if var > -5
                          else COLOR_MUTED))
    _set_cell_text(table.cell(1, 3), str(kpis.get("fatores_urbanos", "—")),
                   align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
    _set_cell_text(table.cell(1, 4), str(kpis.get("denuncias", "—")),
                   align=WD_ALIGN_PARAGRAPH.CENTER, size=10)


def _add_distribuicao_tipo(doc: Document, area: dict) -> None:
    _add_heading(doc, "DISTRIBUIÇÃO POR TIPO DE OCORRÊNCIA", level=2)
    by_day = area.get("temporal", {}).get("byDay") or []
    totals = {"Roubo": 0, "Furto": 0}
    for d in by_day:
        totals["Roubo"] += d.get("Roubo", 0)
        totals["Furto"] += d.get("Furto", 0)
    total = totals["Roubo"] + totals["Furto"]
    if total == 0:
        _add_para(doc, "Sem ocorrências classificadas no período.", italic=True,
                  color=COLOR_MUTED)
        return
    rows = [
        ("Roubo", totals["Roubo"], f"{100 * totals['Roubo'] / total:.1f}%"),
        ("Furto", totals["Furto"], f"{100 * totals['Furto'] / total:.1f}%"),
    ]
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    _set_table_borders(table)
    _header_row(table, ["Tipo", "Quantidade (30d)", "% do total"])
    for i, (t, n, p) in enumerate(rows, start=1):
        _set_cell_text(table.cell(i, 0), t, bold=True, size=10)
        _set_cell_text(table.cell(i, 1), str(n),
                       align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        _set_cell_text(table.cell(i, 2), p,
                       align=WD_ALIGN_PARAGRAPH.CENTER, size=10)


def _add_analise_temporal(doc: Document, area: dict) -> None:
    _add_heading(doc, "ANÁLISE TEMPORAL", level=2)
    _add_para(doc,
              f"Pico identificado em {area.get('peakHours', '—')}, "
              f"principalmente em {area.get('peakDays', '—')}.",
              size=10)

    # Heatmap dia × hora com total (Roubo + Furto)
    by_hour = area.get("temporal", {}).get("byHour") or []
    if not by_hour:
        return
    # Header: dias da semana
    days = ["Dom", "Seg", "Ter", "Qua", "Qui", "Sex", "Sáb"]
    by_day_map = {d.get("day"): d for d in (area.get("temporal", {}).get("byDay") or [])}
    table = doc.add_table(rows=2, cols=len(days) + 1)
    _set_table_borders(table)
    _header_row(table, ["", *days])
    _subheader_cell(table.cell(1, 0), "Total")
    for i, d in enumerate(days, start=1):
        v = by_day_map.get(d, {})
        total = (v.get("Roubo", 0) or 0) + (v.get("Furto", 0) or 0)
        c = table.cell(1, i)
        _set_cell_text(c, str(total),
                       bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        # Sombreamento gradiente — só pra dar pista visual
        if total >= 5:
            _set_cell_shading(c, SHADE_CRITICAL)
        elif total >= 3:
            _set_cell_shading(c, SHADE_HIGH)
        elif total >= 1:
            _set_cell_shading(c, SHADE_MEDIUM)

    # Linha resumo por faixa horária
    _add_para(doc, "Distribuição por hora do dia (Roubo + Furto):",
              size=10)
    table = doc.add_table(rows=2, cols=len(by_hour))
    _set_table_borders(table)
    for i, h in enumerate(by_hour):
        c0 = table.cell(0, i)
        _set_cell_shading(c0, SHADE_HEADER)
        _set_cell_text(c0, h.get("hour", ""), bold=True,
                       color=RGBColor(0xFF, 0xFF, 0xFF), size=8,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        total = (h.get("Roubo", 0) or 0) + (h.get("Furto", 0) or 0)
        c1 = table.cell(1, i)
        _set_cell_text(c1, str(total), size=8,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        if total >= 5:
            _set_cell_shading(c1, SHADE_CRITICAL)
        elif total >= 3:
            _set_cell_shading(c1, SHADE_HIGH)
        elif total >= 1:
            _set_cell_shading(c1, SHADE_MEDIUM)


def _add_dinamica(doc: Document, area: dict) -> None:
    _add_heading(doc, "2. DINÂMICA CRIMINAL", level=1)
    dyn = area.get("dynamics") or {}
    sources = dyn.get("sources") or {}
    _add_para(doc,
              f"Síntese gerada por IA sobre {sources.get('relints', 0)} RELINT(s), "
              f"{sources.get('denuncias', 0)} denúncia(s) do Disque Denúncia e "
              f"{sources.get('ocorrencias', 0)} ocorrência(s) registradas no período.",
              italic=True, color=COLOR_MUTED, size=9)

    rows = [
        ("Modus Operandi", dyn.get("modusOperandi", "—")),
        ("Perfil de suspeitos", dyn.get("suspectProfile", "—")),
        ("Rotas de fuga e dispersão", dyn.get("escapeRoutes", "—")),
        ("Pontos de receptação", dyn.get("receivingPoints", "—")),
    ]
    table = doc.add_table(rows=len(rows), cols=2)
    _set_table_borders(table)
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(12.5)
    for i, (k, v) in enumerate(rows):
        _subheader_cell(table.cell(i, 0), k)
        _set_cell_text(table.cell(i, 1), v, size=10)


def _add_efetivo(doc: Document, area: dict) -> None:
    _add_heading(doc, "3. EFETIVO EMPREGADO — FORÇA MUNICIPAL", level=1)
    # Inferência simples — total 600 agentes para 9 áreas; ajustar por risco
    risk = area.get("risk", "medium")
    risk_factor = {"critical": 1.4, "high": 1.15, "medium": 1.0, "low": 0.8}.get(risk, 1.0)
    base = round(600 / 9 * risk_factor)
    suggested_modal = "motorizada + a pé" if risk in ("critical", "high") else "a pé"

    coincidences = area.get("coincidences") or []
    suggested_locations = ", ".join(
        c.get("location", "").replace("Hotspot ", "Hotspot ")
        for c in coincidences[:3]
    ) or "—"
    suggested_hours = area.get("peakHours", "—")
    suggested_days = area.get("peakDays", "—")

    rows = [
        ("Nº de agentes por turno", f"{base} (de 600 totais distribuídos nas 9 áreas)",
         "Calibrado pelo risco da área"),
        ("Locais de cobertura prioritários", suggested_locations,
         "Top 3 coincidências por score de risco"),
        ("Horário de cobertura recomendado", suggested_hours,
         "Pico horário identificado pelos dados"),
        ("Dias de cobertura recomendados", suggested_days,
         "Dias com maior incidência no período"),
        ("Modalidade de emprego", suggested_modal,
         "Ajustada pela dinâmica criminal e nível de risco"),
    ]
    table = doc.add_table(rows=len(rows) + 1, cols=3)
    _set_table_borders(table)
    _header_row(table, ["Campo", "Sugestão da plataforma", "Justificativa"])
    table.columns[0].width = Cm(4.5)
    table.columns[1].width = Cm(7.0)
    table.columns[2].width = Cm(5.5)
    for i, (k, v, j) in enumerate(rows, start=1):
        _set_cell_text(table.cell(i, 0), k, bold=True, size=10)
        _set_cell_text(table.cell(i, 1), v, size=10)
        _set_cell_text(table.cell(i, 2), j, size=9, color=COLOR_MUTED)


def _add_fatores(doc: Document, area: dict) -> None:
    _add_heading(doc, "4. FATORES DE INCIDÊNCIA CRIMINAL", level=1)
    factors = area.get("urbanFactors") or []
    if not factors:
        _add_para(doc, "Nenhum fator urbano ativo registrado na área.",
                  italic=True, color=COLOR_MUTED)
        return
    # Agrupa: type → (count, orgao)
    by_type: dict[str, dict[str, Any]] = {}
    for f in factors:
        t = f.get("type", "—")
        if t not in by_type:
            by_type[t] = {"count": 0, "orgao": f.get("orgao", "—"),
                          "category": f.get("category", "—")}
        by_type[t]["count"] += 1
    # Ordena por contagem decrescente
    items = sorted(by_type.items(), key=lambda x: -x[1]["count"])

    table = doc.add_table(rows=len(items) + 1, cols=4)
    _set_table_borders(table)
    _header_row(table, ["Fator identificado", "Categoria",
                        "Ocorrências no polígono", "Responsável"])
    table.columns[0].width = Cm(7.0)
    table.columns[1].width = Cm(3.5)
    table.columns[2].width = Cm(2.5)
    table.columns[3].width = Cm(4.0)
    for i, (t, meta) in enumerate(items, start=1):
        _set_cell_text(table.cell(i, 0), t, size=10)
        _set_cell_text(table.cell(i, 1), meta["category"], size=9,
                       color=COLOR_MUTED)
        _set_cell_text(table.cell(i, 2), str(meta["count"]),
                       bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        _set_cell_text(table.cell(i, 3), meta["orgao"], bold=True, size=10)
        if i % 2 == 0:
            for c in table.rows[i].cells:
                _set_cell_shading(c, SHADE_ROW_ALT)

    _add_para(doc,
              f"Total de {len(factors)} fatores ativos · "
              f"{len(items)} tipos distintos · "
              f"{len({m['orgao'] for m in by_type.values()})} órgão(s) responsável(eis).",
              italic=True, color=COLOR_MUTED, size=9)


def _add_coincidencias(doc: Document, area: dict) -> None:
    _add_heading(doc, "PAINEL DE COINCIDÊNCIAS", level=2)
    coins = area.get("coincidences") or []
    if not coins:
        _add_para(doc, "Nenhuma coincidência de alto risco identificada.",
                  italic=True, color=COLOR_MUTED)
        return
    table = doc.add_table(rows=len(coins) + 1, cols=5)
    _set_table_borders(table)
    _header_row(table, ["ID", "Local", "Mancha criminal",
                        "Fator + janela horária", "Risco"])
    for i, c in enumerate(coins, start=1):
        _set_cell_text(table.cell(i, 0), c.get("id", ""), bold=True, size=9)
        _set_cell_text(table.cell(i, 1), c.get("location", ""), size=9)
        _set_cell_text(table.cell(i, 2), c.get("crime", ""), size=9)
        _set_cell_text(table.cell(i, 3),
                       f"{c.get('factor', '')} · {c.get('timeWindow', '')} · "
                       f"{c.get('operationalGap', '')}", size=9)
        risk = c.get("risk", 0)
        _set_cell_text(table.cell(i, 4), str(risk),
                       bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        if risk >= 85:
            _set_cell_shading(table.cell(i, 4), SHADE_CRITICAL)
        elif risk >= 75:
            _set_cell_shading(table.cell(i, 4), SHADE_HIGH)
        else:
            _set_cell_shading(table.cell(i, 4), SHADE_MEDIUM)


def _add_plano_acao(doc: Document, area: dict) -> None:
    _add_heading(doc, "5. PLANO DE AÇÃO E RESPONSABILIZAÇÃO", level=1)
    _add_para(doc,
              "Plano gerado pela IA com base nas coincidências críticas e nos fatores "
              "urbanos do polígono. Deve ser revisado e formalizado na reunião CompStat.",
              italic=True, color=COLOR_MUTED, size=9)
    actions = area.get("actionPlan") or []
    if not actions:
        _add_para(doc, "Nenhuma ação sugerida.", italic=True, color=COLOR_MUTED)
        return
    table = doc.add_table(rows=len(actions) + 1, cols=4)
    _set_table_borders(table)
    _header_row(table, ["Responsável", "Ação", "Justificativa", "Prioridade"])
    table.columns[0].width = Cm(2.5)
    table.columns[1].width = Cm(6.5)
    table.columns[2].width = Cm(6.0)
    table.columns[3].width = Cm(2.0)
    for i, a in enumerate(actions, start=1):
        _set_cell_text(table.cell(i, 0), a.get("responsible", ""), bold=True, size=10)
        _set_cell_text(table.cell(i, 1), a.get("action", ""), size=10)
        _set_cell_text(table.cell(i, 2), a.get("justification", ""), size=9,
                       color=COLOR_MUTED)
        prio = (a.get("priority") or "").lower()
        prio_label = {"alta": "Alta", "media": "Média", "baixa": "Baixa"}.get(prio, "—")
        c_prio = table.cell(i, 3)
        _set_cell_text(c_prio, prio_label, bold=True,
                       align=WD_ALIGN_PARAGRAPH.CENTER, size=10)
        if prio == "alta":
            _set_cell_shading(c_prio, SHADE_CRITICAL)
        elif prio == "media":
            _set_cell_shading(c_prio, SHADE_MEDIUM)


def _add_footer_note(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run(
        "Relatório gerado automaticamente pelo CompStat Rio — Plataforma de "
        "Inteligência Criminal · Secretaria-Geral do CompStat Municipal."
    )
    run.italic = True
    run.font.size = Pt(8)
    _set_run_color(run, COLOR_MUTED)


# ---------------------------------------------------------------------------
# Public entrypoint
# ---------------------------------------------------------------------------

def build_report(area: dict, reference_date: str | None = None) -> bytes:
    """Monta o .docx completo e retorna os bytes serializados."""
    doc = Document()
    # Margens mais apertadas para caber as tabelas largas
    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.6)
        section.left_margin = Cm(1.4)
        section.right_margin = Cm(1.4)

    _add_cover(doc, area, reference_date)
    _add_map_image(doc, area)
    _add_resumo_executivo(doc, area)
    _add_identificacao(doc, area)
    _add_indicadores(doc, area)
    _add_distribuicao_tipo(doc, area)
    _add_analise_temporal(doc, area)
    _add_dinamica(doc, area)
    _add_efetivo(doc, area)
    _add_fatores(doc, area)
    _add_coincidencias(doc, area)
    _add_plano_acao(doc, area)
    _add_footer_note(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
